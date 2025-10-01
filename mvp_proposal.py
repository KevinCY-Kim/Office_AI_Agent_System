# ===============================
# 0. 라이브러리 임포트
# ===============================
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline

# ===== RAG 유틸 =====
import json
from sentence_transformers import SentenceTransformer
import faiss
from rank_bm25 import BM25Okapi

# ===============================
# 1. 인덱서 / 검색기
# ===============================
class TextIndexer:
    def __init__(self, model_name="intfloat/multilingual-e5-base"):
        self.model = SentenceTransformer(model_name)
        self.idx = None
        self.chunks = []

    def build(self):
        texts = [c["text"] for c in self.chunks]
        if not texts:
            self.idx = None
            return
        embs = self.model.encode(texts, normalize_embeddings=True)
        dim = embs.shape[1]
        self.idx = faiss.IndexFlatIP(dim)
        self.idx.add(embs.astype("float32"))

    def search_dense(self, query: str, topk=4):
        if self.idx is None or not self.chunks:
            return []
        qv = self.model.encode([query], normalize_embeddings=True)
        D, I = self.idx.search(qv, topk)
        items = []
        for i in I[0]:
            if i < 0:
                continue
            items.append(self.chunks[i])
        return items

class HybridRetriever:
    def __init__(self, dense_indexer):
        self.dense = dense_indexer
        self.corpus = [c["text"] for c in self.dense.chunks]
        self.bm25 = BM25Okapi([t.split() for t in self.corpus]) if self.corpus else None

    def search(self, query, k_dense=4, k_bm25=4):
        dense_hits = self.dense.search_dense(query, topk=k_dense)
        bm_hits = []
        if self.bm25 is not None:
            bm_ids = self.bm25.get_top_n(query.split(), list(range(len(self.corpus))), n=k_bm25)
            bm_hits = [self.dense.chunks[i] for i in bm_ids]
        seen, results = set(), []
        for it in dense_hits + bm_hits:
            key = it["text"][:80]
            if key in seen: 
                continue
            seen.add(key)
            results.append(it)
        return results[:max(k_dense, k_bm25)]

# ===============================
# 2. 모델 로드
# ===============================
model_name = "skt/A.X-4.0-Light"
tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModelForCausalLM.from_pretrained(
    model_name,
    device_map="auto",
    torch_dtype="auto"
)
generator = pipeline("text-generation", model=model, tokenizer=tokenizer)

# ===============================
# 3. rag_chunks.json 불러오기 → RAG 인덱스 생성
# ===============================
rag_chunks_path = "/home/alpaco/kimcy/Office_AI_Agent_System/report/rag_chunks.json"
with open(rag_chunks_path, "r", encoding="utf-8") as f:
    rag_chunks_list = json.load(f)

indexer = TextIndexer()
indexer.chunks = [{"title": f"chunk_{i+1}", "text": txt} for i, txt in enumerate(rag_chunks_list)]
indexer.build()
retriever = HybridRetriever(indexer)
print(f"총 {len(indexer.chunks)}개의 청크로 RAG 인덱스 생성 완료")

# ===============================
# 3-1. R&D 가이드라인 불러오기
# ===============================
guideline_path = "/home/alpaco/kimcy/Office_AI_Agent_System/config/rnd_guidelines/rnd_guidline.json"
with open(guideline_path, "r", encoding="utf-8") as f:
    rnd_rules = json.load(f)

def flatten_guidelines(rnd_rules):
    all_rules = []
    for sec in rnd_rules.get("sections", []):
        for item in sec.get("items", []):
            if isinstance(item, str):
                all_rules.append(item)
            elif isinstance(item, dict):
                if "content" in item:
                    all_rules.append(item["content"])
                if "examples" in item:
                    all_rules.append(f"(예시: {item['examples']})")
                if "reasons" in item:
                    all_rules.append(f"(이유: {item['reasons']})")
    return all_rules

guideline_list = flatten_guidelines(rnd_rules)

# ===============================
# 4. 사용자 입력
# ===============================
project_name = input("사업명: ")
company_name = input("회사명: ")
manager_name = input("담당자: ")
keywords = input("문제 정의 키워드 (쉼표로 구분): ")
budget = input("예산(단위: 백만원): ")

# ===============================
# 5. 섹션별 프롬프트
# ===============================
section_roles = {
    "연구개발 목표": "당신은 R&D PMO입니다. 단계/일괄 협약의 최종목표를 500자 내외로 명확·간결·정량화하여 작성합니다. 핵심 성능지표(KPI), 달성 기준(수치/단위/마일스톤), 검증 방법을 포함하고 모호한 표현은 배제합니다.",
    "연구개발 내용": "당신은 기술 총괄(Tech Lead)입니다. 전체 연구범위를 1,000자 내외로 구조화해 기술 요소, 서브태스크, 인터페이스, 데이터/시스템 흐름을 설명하고 표준·규격·평가계획을 명시합니다.",
    "연구개발성과 활용계획 및 기대효과": "당신은 사업전략/사업개발(BD) 담당자입니다. 수요처, 적용 시나리오, 도입·확산 경로, 수익/비용 구조, 경제적 파급효과를 500자 내외로 정량·정성 지표와 함께 제시합니다.",
    "연구기획과제의 개요": "당신은 제안서 총괄 에디터입니다. 목적·필요성·기대효과를 일관된 논리로 요약해 과제가 해결하는 문제와 중요성을 한눈에 보이게 작성합니다.",
    "연구개발과제의 배경": "당신은 정책/RFP 적합성 분석가입니다. 관련 선행연구·시장/기술 동향·정부 정책·RFP/품목요약서 부합성을 근거와 함께 정리하고 제안 맥락을 명확히 합니다.",
    "연구개발과제의 필요성": "당신은 산업분석가입니다. 현황·문제점·시장규모/성장률·규제 및 정책 요구를 데이터로 제시하고, 해결 필요성을 인과적으로 설득력 있게 제시합니다.",
    "보안등급의 분류 및 해당 사유": "당신은 보안관리 책임자입니다. 국가연구개발혁신법 시행령 제45조 및 산업기술혁신사업 보안관리요령 제9조 기준을 근거로 보안등급과 결정 사유를 간결히 기재합니다.",
    "기술개발 핵심어(키워드)": "당신은 표준/용어 관리자입니다. 과제의 핵심 용어 5개를 한글/영문 정식 명칭으로 제시하고, 표준(협회/학회) 정의에 부합하도록 작성합니다.",
    "연차별 개발목표": "당신은 일정/성과관리 PM입니다. 연차별(1년차~n년차) 목표를 기관(주관/공동/참여 연구원)별로 구분해 KPI·마일스톤·검증기준을 정량화하여 제시합니다.",
    "연차별 개발내용 및 범위": "당신은 공동연구 컨소시엄 코디네이터입니다. 기관별 역할·범위·인계·의존성을 명확히 기술하고 중복/누락 없이 연차별 산출물과 책임을 표로 정리합니다(공동기관 없으면 생략).",
    "추진방법 및 전략": "당신은 기술전략/실험 설계 책임자입니다. 방법론(데이터·알고리즘·장비), 리스크와 대응책, 실험/검증 계획(평가지표·샘플수·통계/검증 절차)을 구체적으로 기술합니다.",
    "과제 성과의 활용방안": "당신은 제품/사업화 매니저입니다. 성과의 적용 분야, 기술 파급효과, 에너지 절감·환경 개선 등 기술적·사회적 효익을 사용 시나리오와 함께 제시합니다.",
    "신규사업 신설의 기대효과": "당신은 전략기획 임원입니다. 시장 창출, 일자리, 수입대체, 수출 증대, 비용 절감 등 경제·산업적 효과를 정량 지표(금액, 비율, 기간)와 함께 제시합니다.",
    "사회적 가치 창출 계획": "당신은 ESG/사회가치 책임자입니다. 개요-비전-목표-세부계획-기대효과 체계로 13개 사회적 가치 범주와의 연계를 명확히 하고 측정 가능한 지표를 포함합니다.",
    "사회적 가치창출의 기대효과": "당신은 임팩트 평가자입니다. 보건·안전·포용·지역·환경·민주성 등 사회적 가치 지표를 중심으로 성과/파급효과를 정량·정성으로 제시합니다.",
    "경제적 성과창출의 기대효과": "당신은 재무 담당자입니다(기업 작성). 매출/원가/영업이익, ROI/NPV, 고용효과 등 재무적 성과 전망을 가정과 산식(간단) 포함하여 명료하게 제시합니다.",
    "신규 인력 채용 계획 및 활용 방안": "당신은 HR 책임자입니다. 신규/기존 채용 구분, 채용 시점·역할·배치·활용 계획, 역량 매핑과 교육/온보딩 계획을 일정표와 함께 제시합니다."
    }

section_queries = { 
    "연구개발 목표":"최종목표(단계/일괄 협약목표)를 과제의 연구기획목표를 500자 내외로 기재합니다.",
    "연구개발 내용": "전체내용을 1,000자 내외로 기재합니다.",
    "연구개발성과 활용계획 및 기대효과": "연구기획의 수요처, 활용내용, 경제적 파급효과 등을 500자 내외로 기재합니다(연구시설ㆍ장비 구축을 목적으로 하는 과제의 경우에 연구시설ㆍ장비를 활용한 성과관리 및 자립운영계획, 수입금 관리 및 운영계획 등).",
    "연구기획과제의 개요": "연구기획과제의 개요는 연구개발과제의 목적, 필요성, 기대효과 등을 종합적으로 고려하여 작성합니다. 이를 통해 연구개발과제가 어떤 문제를 해결하고자 하는지, 왜 중요한지 명확히 제시합니다.",
    "연구개발과제의 배경": "구개발과제와 관련되는 연구개발과제의 배경 및 필요성, 정부 정책 및 RFP/품목요약서의 부합성 등을 종합적으로 기재합니다. 이를 통해 연구개발과제가 어떤 맥락에서 제안되었는지, 어떤 문제를 해결하고자 하는지 명확히 제시합니다.",
    "연구개발과제의 필요성":"연구개발과제의 필요성은 해당 기술 및 산업의 현황, 문제점, 시장 동향, 정책적 요구사항 등을 종합적으로 고려하여 작성합니다. 이를 통해 연구개발과제가 왜 필요한지, 어떤 문제를 해결하고자 하는지 명확히 제시합니다.",
    "보안등급의 분류 및 해당 사유": "국가연구개발혁신법 시행령 제45조(연구개발과제에 대한 보안과제의 분류) 및 산업기술혁신사업 보안관리요령 제9조(보안등급 분류 기준)을 참조하여, 계획서 표지에 있는 보안등급 분류에 대한 결정사유 기입합니다.",
    "기술개발 핵심어(키워드)": "핵심어는 동일 개발과제의 핵심적 용어로써 과제 관련 특수 용어로써 관련 업계, 협회나 학회 등에서 표준화되어 정의되었거나 일반화된 정식 명칭을 기재하며, 5개 단어를 한글 및 영문으로 반드시 기입하여야 함",
    "연차별 개발목표" : "1년차도, 2년차도, n년차도 각각의 연차별 개발목표를 주관연구개발기관, 공동연구개발기관, 참여연구원별로 구분하여 작성합니다.",
    "연차별 개발내용 및 범위" : "주관연구개발기관 및 공동연구개발기관이 담당하는 부분을 기술․표시하고, 연구개발기관별 연차별 개발목표, 내용 및 범위가 명확히 드러나도록 기술(공동연구개발기관이 없는 경우 생략)합니다.",
    "추진방법 및 전략" : "개발목표 달성을 위하여 무엇을 활용하고 어떻게 수행할 것인지 등 수행 방법을 구체적으로 기술하고, 세부개발 내용별 수행 방법, 수행 과정 중 예측되는 장애 요소 및 그 해결 방안, 계획된 실험과정 등을 기술합니다.",
    "과제 성과의 활용방안" : "연구개발과제 수행에 따라 예상되는 성과와 그 활용분야 및 활용방안을 기재하고, 기술적 측면은 해당 기술의 향상, 다른 기술로의 파급 효과 및 기술개발에 따른 에너지 절약 또는 환경 개선 효과 등을 서술합니다.",
    "신규사업 신설의 기대효과" : "연구개발성과의 과학ㆍ기술적, 경제ㆍ산업적, 사회적 측면에서 기대효과ㆍ파급효과 등을 기재하고, 경제․산업적 측면에는 시장 창출 및 일자리 창출 효과, 수입 대체 효과, 수출 증대 효과, 비용 절감 등의 경제적 효과와 산업발전에의 영향 등 산업적 효과를 서술합니다.",
    "사회적 가치 창출 계획" : "개요, 추진전략(비전), 추진목표, 세부계획, 기대효과로 구분하여 작성합니다.사회적 가치란 사회적·경제적·환경적·문화적 영역에서 공공의 이익과 공동체 발전에 기여하는 가치로서 13가지(인간의 존엄성을 유지하는 기본권리로서 인권보호, 재난과 사고로부터 안전한 근로 생활환경의 유지, 건강한 생활이 가능한 보건복지의 제공, 노동권의 보장과 근로조건의 향상, 사회적 약자에 대한 기회제공과 사회통합, 대기업·중소기업간 상생과 협력, 품위있는 삶을 누릴 수 있는 양질의 일자리 창출, 지역사회 활성화와 공동체 복원, 경제활동을 통한 이익이 지역에 순환되는 지역경제 공헌, 윤리적 생산과 유통을 포함한 기업의 자발적인 사회적 책임 이행, 환경의 지속가능성 보전, 시민적 권리로서 민주적 의사결정과 참여의 실현, 그 밖에 공동체의 이익 실현과 공공성 강화)을 포괄하는 가치를 의미합니다.",
    "사회적 가치창출의 기대효과": "연구개발과제 수행에 따라 예상되는 성과와 그 활용분야 및 활용방안을 기재하고, 기술적 측면은 해당 기술의 향상, 다른 기술로의 파급 효과 및 기술개발에 따른 에너지 절약 또는 환경 개선 효과 등을 서술합니다.",
    "경제적 성과창출의 기대효과" : "주관/공동연구개발기관 중 기업만 작성합니다.",
    "신규 인력 채용 계획 및 활용 방안" : "신규 채용 여부는 신규 채용인 경우와 기존인 경우로 표기합니다. 또한, 신규 채용 구분 여부는 동 과제 수행을 위해 사업 공고일 기준 6개월 이전에 신규로 채용했거나 과제 전체 연구개발기간 중 채용 계획이 있는 경우로 구분하여 서술합니다.",
    }

# ===============================
# 6. 자동 문장 생성 함수
# ===============================
def search_contexts(section: str, topk=4):
    query = section_queries.get(section, section)
    hits = retriever.search(query, k_dense=4, k_bm25=4)
    contexts, max_ctx_len = [], 900
    for h in hits:
        t = h["text"]
        if len(t) > max_ctx_len:
            t = t[:max_ctx_len] + "..."
        contexts.append(t)
    return contexts

def build_prompt_with_context(section, role_instruction, contexts, guideline_list=None):
    ctx_block = "\n\n".join([f"[근거]\n{c}" for c in contexts]) if contexts else "[근거]\n(해당 섹션에 대한 근거 없음)"
    guideline_block = ""
    if guideline_list:
        guideline_block = "\n".join([f"- {rule}" for rule in guideline_list])
    return f"""
역할: {role_instruction}
작성 항목: [{section}]
회사명: {company_name}
사업명: {project_name}

작성 조건:
- 아래 근거 텍스트 안에서만 답변할 것
- 출처 밖 가정 금지, 누락 정보는 '확인 필요'로 표기
- 수치/날짜/명칭은 원문 그대로 유지
- 국가사업 제안서 톤
- 5 ~ 8문장
- 다른 항목과 중복 표현 최소화
- 반드시 아래 R&D 가이드라인 준수

R&D 가이드라인:
{guideline_block}

{ctx_block}
""".strip()

def generate_text(section, keywords=""):
    role_instruction = section_roles.get(section, "")
    contexts = search_contexts(section, topk=4)
    prompt = build_prompt_with_context(section, role_instruction, contexts, guideline_list)
    # output = generator(
    #     prompt,
    #     max_new_tokens=220,
    #     do_sample=False,
    #     repetition_penalty=1.05,
    #     eos_token_id=tokenizer.eos_token_id
    # )
    output = generator(
    prompt,
    max_new_tokens=5000,          # 원하는 생성 길이
    do_sample=True,                # 긴 텍스트엔 샘플링을 권장
    temperature=0.7,               # 다양성 조절
    repetition_penalty=1.2,        # 중복 방지
    eos_token_id=tokenizer.eos_token_id
)
    text = output[0]["generated_text"]
    return text[len(prompt):].strip() if text.startswith(prompt) else text.strip()

# ===============================
# 7. DOCX 생성
# ===============================
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
font.size = Pt(11)

doc.add_heading("국가사업 제안서 (간략 버전)", 0)
doc.add_paragraph(f"사업명: {project_name}")
doc.add_paragraph(f"회사명: {company_name}")
doc.add_paragraph(f"담당자: {manager_name}")
doc.add_paragraph(f"예산: {budget} 백만원")
doc.add_paragraph("")

sections = ["사업 개요", "문제 정의", "해결 방안", "기대 효과"]

for section in sections:
    doc.add_heading(section, level=1)
    text = generate_text(section, keywords)
    doc.add_paragraph(text)

doc.add_page_break()

# ===============================
# 8. 파일 저장
# ===============================
output_file = "mvp_proposal_ai.docx"
doc.save(output_file)
print(f"✅ '{output_file}' 파일이 생성되었습니다!")
