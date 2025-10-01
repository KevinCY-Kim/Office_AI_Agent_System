"""
테스트 실행용 최소 코드 (RAG 기반 국가사업 제안서 자동 생성)

✅ 실행에 필요한 부분만 남김
✅ 이미 마무리된 단계(청크 생성 등)는 주석 처리 → 필요 시 참고
"""

# ===============================
# 0. 라이브러리 임포트
# ===============================
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline

# ===== RAG 유틸 =====
import json
from sentence_transformers import SentenceTransformer
import faiss
from rank_bm25 import BM25Okapi

# ===============================
# 1. 인덱서 / 검색기
# ===============================
class TextIndexer:
    def __init__(self, model_name="intfloat/multilingual-e5-base"):
        self.model = SentenceTransformer(model_name)
        self.idx = None
        self.chunks = []

    def build(self):
        texts = [c["text"] for c in self.chunks]
        if not texts:
            self.idx = None
            return
        embs = self.model.encode(texts, normalize_embeddings=True)
        dim = embs.shape[1]
        self.idx = faiss.IndexFlatIP(dim)
        self.idx.add(embs.astype("float32"))

    def search_dense(self, query: str, topk=4):
        if self.idx is None or not self.chunks:
            return []
        qv = self.model.encode([query], normalize_embeddings=True)
        D, I = self.idx.search(qv, topk)
        items = []
        for i in I[0]:
            if i < 0:
                continue
            items.append(self.chunks[i])
        return items

class HybridRetriever:
    def __init__(self, dense_indexer):
        self.dense = dense_indexer
        self.corpus = [c["text"] for c in self.dense.chunks]
        self.bm25 = BM25Okapi([t.split() for t in self.corpus]) if self.corpus else None

    def search(self, query, k_dense=4, k_bm25=4):
        dense_hits = self.dense.search_dense(query, topk=k_dense)
        bm_hits = []
        if self.bm25 is not None:
            bm_ids = self.bm25.get_top_n(query.split(), list(range(len(self.corpus))), n=k_bm25)
            bm_hits = [self.dense.chunks[i] for i in bm_ids]
        seen, results = set(), []
        for it in dense_hits + bm_hits:
            key = it["text"][:80]
            if key in seen: 
                continue
            seen.add(key)
            results.append(it)
        return results[:max(k_dense, k_bm25)]

# ===============================
# 2. 모델 로드
# ===============================
model_name = "skt/A.X-4.0-Light"
tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModelForCausalLM.from_pretrained(
    model_name,
    device_map="auto",
    torch_dtype="auto"
)
generator = pipeline("text-generation", model=model, tokenizer=tokenizer)

# ===============================
# 3. rag_chunks.json 불러오기 → RAG 인덱스 생성
# ===============================
rag_chunks_path = "/home/alpaco/kimcy/Office_AI_Agent_System/report/rag_chunks.json"
with open(rag_chunks_path, "r", encoding="utf-8") as f:
    rag_chunks_list = json.load(f)

indexer = TextIndexer()
indexer.chunks = [{"title": f"chunk_{i+1}", "text": txt} for i, txt in enumerate(rag_chunks_list)]
indexer.build()
retriever = HybridRetriever(indexer)
print(f"총 {len(indexer.chunks)}개의 청크로 RAG 인덱스 생성 완료")

# ===============================
# 4. 사용자 입력
# ===============================
project_name = input("사업명: ")
company_name = input("회사명: ")
manager_name = input("담당자: ")
keywords = input("문제 정의 키워드 (쉼표로 구분): ")
budget = input("예산(단위: 백만원): ")

# ===============================
# 5. 섹션별 프롬프트
# ===============================
section_roles = {
    "사업 개요": "당신은 투자 심사역입니다. 투자자의 관점에서 사업 개요를 설명하세요.",
    "문제 정의": "당신은 현업 부서장입니다. 실제 업무에서 느끼는 문제 정의를 강조하세요.",
    "해결 방안": "당신은 해당 분야의 전문적인 분석가입니다. 기술적 해결 방안을 구체적으로 제시하세요.",
    "기대 효과": "당신은 CEO입니다. 경영적 기대 효과를 전략적 가치 중심으로 설명하세요."
}

section_queries = {
    "사업 개요": "회사 개요 설립일 본점 소재지 주된 사업 요약",
    "문제 정의": f"{keywords} 관련 당사 사업환경 리스크 또는 문제 서술",
    "해결 방안": f"{keywords} 해결 기술/제품/서비스 방향 핵심 근거",
    "기대 효과": "경영성과, 시장성, 경쟁우위, 수익성, 전략적 기대효과"
}

# ===============================
# 6. 자동 문장 생성 함수
# ===============================
def search_contexts(section: str, topk=4):
    query = section_queries.get(section, section)
    hits = retriever.search(query, k_dense=4, k_bm25=4)
    contexts, max_ctx_len = [], 900
    for h in hits:
        t = h["text"]
        if len(t) > max_ctx_len:
            t = t[:max_ctx_len] + "..."
        contexts.append(t)
    return contexts

def build_prompt_with_context(section, role_instruction, contexts):
    ctx_block = "\n\n".join([f"[근거]\n{c}" for c in contexts]) if contexts else "[근거]\n(해당 섹션에 대한 근거 없음)"
    return f"""
역할: {role_instruction}
작성 항목: [{section}]
회사명: {company_name}
사업명: {project_name}

작성 조건:
- 아래 근거 텍스트 안에서만 답변할 것
- 출처 밖 가정 금지, 누락 정보는 '확인 필요'로 표기
- 수치/날짜/명칭은 원문 그대로 유지
- 국가사업 제안서 톤
- 3~5문장
- 다른 항목과 중복 표현 최소화

{ctx_block}
""".strip()

def generate_text(section, keywords=""):
    role_instruction = section_roles.get(section, "")
    contexts = search_contexts(section, topk=4)
    prompt = build_prompt_with_context(section, role_instruction, contexts)
    output = generator(
        prompt,
        max_new_tokens=220,
        do_sample=False,
        repetition_penalty=1.05,
        eos_token_id=tokenizer.eos_token_id
    )
    text = output[0]["generated_text"]
    return text[len(prompt):].strip() if text.startswith(prompt) else text.strip()

# ===============================
# 7. DOCX 생성
# ===============================
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
font.size = Pt(11)

doc.add_heading("국가사업 제안서 (간략 버전)", 0)
doc.add_paragraph(f"사업명: {project_name}")
doc.add_paragraph(f"회사명: {company_name}")
doc.add_paragraph(f"담당자: {manager_name}")
doc.add_paragraph(f"예산: {budget} 백만원")
doc.add_paragraph("")

sections = ["사업 개요", "문제 정의", "해결 방안", "기대 효과"]

for section in sections:
    doc.add_heading(section, level=1)
    doc.add_paragraph(generate_text(section, keywords))

doc.add_page_break()

# ===============================
# 8. 파일 저장
# ===============================
output_file = "mvp_proposal_ai.docx"
doc.save(output_file)
print(f"✅ '{output_file}' 파일이 생성되었습니다!")


"""
# ===============================
# (참고용) 청크 생성 코드
# ===============================
# 아래 코드는 최초 1회 실행용. 이미 rag_chunks.json 생성 후엔 불필요.

# from langchain_community.document_loaders import PyPDFLoader
# from langchain_text_splitters import RecursiveCharacterTextSplitter

# loader = PyPDFLoader("input.pdf")
# docs = loader.load()

# text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
# chunks = text_splitter.split_documents(docs)

# with open("rag_chunks.json", "w", encoding="utf-8") as f:
#     json.dump([c.page_content for c in chunks], f, ensure_ascii=False, indent=2)
"""
