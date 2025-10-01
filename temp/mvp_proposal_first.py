from docx import Document
from transformers import AutoTokenizer, AutoModelForCausalLM, pipeline
from docx.shared import Pt
from docx.oxml.ns import qn

# ===============================
# 1. 모델 로드
# ===============================
model_name = "skt/A.X-4.0-Light"
tokenizer = AutoTokenizer.from_pretrained(model_name)
model = AutoModelForCausalLM.from_pretrained(
    model_name,
    device_map="auto",   # GPU 자동 할당
    torch_dtype="auto"
)

generator = pipeline("text-generation", model=model, tokenizer=tokenizer)

# ===============================
# 2. 사용자 입력 받기
# ===============================
project_name = input("사업명: ")
company_name = input("회사명: ")
manager_name = input("담당자: ")
keywords = input("문제 정의 키워드 (쉼표로 구분): ")
budget = input("예산(단위: 백만원): ")

# ===============================
# 3. 섹션별 롤플레이 프롬프트 정의
# ===============================
section_roles = {
    "사업 개요": "당신은 투자 심사역입니다. 투자자의 관점에서 사업 개요를 설명하세요.",
    "문제 정의": "당신은 현업 부서장입니다. 실제 업무에서 느끼는 문제 정의를 강조하세요.",
    "해결 방안": "당신은 해당 분야의 전문적인 분석가입니다. 기술적 해결 방안을 구체적으로 제시하세요.",
    "기대 효과": "당신은 CEO입니다. 경영적 기대 효과를 전략적 가치 중심으로 설명하세요."
}

# ===============================
# 4. 자동 문장 생성 함수
# ===============================
def generate_text(section, keywords=""):
    role_instruction = section_roles.get(section, "")
    prompt = f"""
    역할: {role_instruction}
    작성 항목: [{section}]
    키워드: {keywords}
    회사명: {company_name}
    사업명: {project_name}
    조건:
    - 국가사업 제안서 형식으로 작성
    - 3~5문장
    - 다른 항목과 중복되는 표현은 피할 것
    - 전문적인 톤 유지
    """
    output = generator(prompt, max_new_tokens=200, do_sample=True, top_p=0.9, temperature=0.7)
    return output[0]["generated_text"].replace(prompt, "").strip()

# ===============================
# 5. DOCX 문서 생성
# ===============================
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
font.size = Pt(11)

doc.add_heading("국가사업 제안서 (간략 버전)", 0)
doc.add_paragraph(f"사업명: {project_name}")
doc.add_paragraph(f"회사명: {company_name}")
doc.add_paragraph(f"담당자: {manager_name}")
doc.add_paragraph(f"예산: {budget} 백만원")
doc.add_paragraph("")

sections = ["사업 개요", "문제 정의", "해결 방안", "기대 효과"]

for section in sections:
    doc.add_heading(section, level=1)
    doc.add_paragraph(generate_text(section, keywords))

doc.add_page_break()

# ===============================
# 6. 파일 저장
# ===============================
output_file = "mvp_proposal_ai.docx"
doc.save(output_file)
print(f"✅ '{output_file}' 파일이 생성되었습니다!")
