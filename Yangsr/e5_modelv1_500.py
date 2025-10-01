import os, json, random, math
import numpy as np
import torch
from tqdm.auto import tqdm
import re # 정규표현식 사용을 위해 추가
from transformers import (
    AutoTokenizer, AutoModel, AutoModelForCausalLM, pipeline
)
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

from datasets import Dataset
from torch.utils.data import DataLoader
from typing import List

# ==============================================================================
# 1. 환경 설정 및 경로 변수
# ==============================================================================
GUIDELINE_FILE = '/home/alpaco/kimcy/Office_AI_Agent_System/config/rnd_guidelines/rnd_guidline.json'
RAG_JSON_FILES = ["/home/alpaco/kimcy/Office_AI_Agent_System/chunking/rag_chunks500_50.json"] 

# 사용 모델 명
E5_NAME = "intfloat/multilingual-e5-base" # e5 모델 (검색 관련 로직은 제거됨)
GEN_NAME = "skt/A.X-4.0-Light"              # 생성 LLM

# 하드웨어 설정
DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
CUDA_ID = 0 if torch.cuda.is_available() else -1    

# e5 모델 관련 파라미터 (검색이 비활성화되어 사용되지 않음)
DOCUMENT_PREFIX = "passage: " 
QUERY_PREFIX    = "query: "   
MAXLEN_DOC   = 512            
MAXLEN_QUERY = 512            
BATCH_SIZE   = 32             

GEN_MODE = "pipeline_batch" 
GEN_MAX_NEW_TOKENS = 5000
GEN_DO_SAMPLE = False           
GEN_TEMPERATURE = None          


# ==============================================================================
# 2. 유틸리티 함수
# ==============================================================================
def l2_normalize(t: torch.Tensor, dim=1):
    """L2 정규화 함수 (현재 코드에서는 사용되지 않음)"""
    return torch.nn.functional.normalize(t, p=2, dim=dim)

def first_n_lines(text: str, n_chars=350):
    """텍스트를 띄어쓰기로 합치고 처음 n_chars만큼 반환 (근거 텍스트 간결화)"""
    t = " ".join(str(text).split())
    return t[:n_chars]

def clean_generated_text(text: str) -> str:
    """
    생성된 텍스트에서 프롬프트 부분('#=========== 출력' 이전)과 지정된 Markdown 포맷을 제거합니다.
    """
    # 1. 프롬프트 헤더 이후의 텍스트만 추출
    output_marker = "#=========== 출력"
    if output_marker in text:
        text = text.split(output_marker, 1)[-1].strip()

    # 2. 불용어 및 Markdown 포맷 제거
    text = re.sub(r'(\*\*(.*?)\*\*|#+\s*|\*\*)', lambda m: m.group(2) if m.group(2) else '', text)
    

    return text.strip()

# ==============================================================================
# 3. 데이터 로드 (FAISS 인덱스 구축 로직 제거)
# ==============================================================================
# 1) 가이드라인 및 RAG 데이터 로드
with open(GUIDELINE_FILE, "r", encoding="utf-8") as f:
    # R&D 가이드라인 데이터 로드 (프롬프트 구성에 사용)
    guidelines = json.load(f)

combined_text = ""
for jf in RAG_JSON_FILES:
    with open(jf, "r", encoding="utf-8") as f:
        parsed = json.load(f)
        # RAG 데이터 JSON 파일을 순회하며 텍스트를 결합 (컨텍스트/근거로 활용)
        if isinstance(parsed, dict) and "text" in parsed:
            combined_text += parsed["text"] + "\n\n"
        else:
            combined_text += str(parsed) + "\n\n"
print("[INFO] RAG chunks loaded. (Only combined_text will be used as context.)")

# (e5 문서 임베딩/문서 원문 로드 및 FAISS 인덱스 구축 로직 모두 제거됨)

# ==============================================================================
# 4. e5 기반 검색(Retrieval) 로직 (모두 제거됨)
# ==============================================================================
# e5 모델, 임베딩 함수 (mean_pooling, ensure_e5_model, embed_query_e5) 및
# dense 검색 함수 (search_dense_by_text)는 FAISS 사용 중단으로 모두 제거됨.

# 5) 섹션별 역할/질의 템플릿 
section_roles = { 
    "연구개발 목표": "당신은 R&D PMO입니다. 단계/일괄 협약의 최종목표를 500자 내외로 명확·간결·정량화하여 작성합니다. 핵심 성능지표(KPI), 달성 기준(수치/단위/마일스톤), 검증 방법을 포함하고 모호한 표현은 배제합니다.",
    "연구개발 내용": "당신은 기술 총괄(Tech Lead)입니다. 전체 연구범위를 1,000자 내외로 구조화해 기술 요소, 서브태스크, 인터페이스, 데이터/시스템 흐름을 설명하고 표준·규격·평가계획을 명시합니다.",
    "연구개발성과 활용계획 및 기대효과": "당신은 사업전략/사업개발(BD) 담당자입니다. 수요처, 적용 시나리오, 도입·확산 경로, 수익/비용 구조, 경제적 파급효과를 500자 내외로 정량·정성 지표와 함께 제시합니다.",
    "연구기획과제의 개요": "당신은 제안서 총괄 에디터입니다. 목적·필요성·기대효과를 일관된 논리로 요약해 과제가 해결하는 문제와 중요성을 한눈에 보이게 작성합니다.",
    "연구개발과제의 배경": "당신은 정책/RFP 적합성 분석가입니다. 관련 선행연구·시장/기술 동향·정부 정책·RFP/품목요약서 부합성을 근거와 함께 정리하고 제안 맥락을 명확히 합니다.",
    "연구개발과제의 필요성": "당신은 산업분석가입니다. 현황·문제점·시장규모/성장률·규제 및 정책 요구를 데이터로 제시하고, 해결 필요성을 인과적으로 설득력 있게 제시합니다.",
    "보안등급의 분류 및 해당 사유": "당신은 보안관리 책임자입니다. 국가연구개발혁신법 시행령 제45조 및 산업기술혁신사업 보안관리요령 제9조 기준을 근거로 보안등급과 결정 사유를 간결히 기재합니다.",
    "기술개발 핵심어(키워드)": "당신은 표준/용어 관리자입니다. 과제의 핵심 용어 5개를 한글/영문 정식 명칭으로 제시하고, 표준(협회/학회) 정의에 부합하도록 작성합니다.",
    "연차별 개발목표": "당신은 일정/성과관리 PM입니다. 연차별(1년차~n년차) 목표를 기관(주관/공동/참여 연구원)별로 구분해 KPI·마일스톤·검증기준을 정량화하여 제시합니다.",
    "연차별 개발내용 및 범위": "당신은 공동연구 컨소시엄 코디네이터입니다. 기관별 역할·범위·인계·의존성을 명확히 기술하고 중복/누락 없이 연차별 산출물과 책임을 표로 정리합니다(공동기관 없으면 생략).",
    "추진방법 및 전략": "당신은 기술전략/실험 설계 책임자입니다. 방법론(데이터·알고리즘·장비), 리스크와 대응책, 실험/검증 계획(평가지표·샘플수·통계/검증 절차)을 구체적으로 기술합니다.",
    "과제 성과의 활용방안": "당신은 제품/사업화 매니저입니다. 성과의 적용 분야, 기술 파급효과, 에너지 절감·환경 개선 등 기술적·사회적 효익을 사용 시나리오와 함께 제시합니다.",
    "신규사업 신설의 기대효과": "당신은 전략기획 임원입니다. 시장 창출, 일자리, 수입대체, 수출 증대, 비용 절감 등 경제·산업적 효과를 정량 지표(금액, 비율, 기간)와 함께 제시합니다.",
    "사회적 가치 창출 계획": "당신은 ESG/사회가치 책임자입니다. 개요-비전-목표-세부계획-기대효과 체계로 13개 사회적 가치 범주와의 연계를 명확히 하고 측정 가능한 지표를 포함합니다.",
    "사회적 가치창출의 기대효과": "당신은 임팩트 평가자입니다. 보건·안전·포용·지역·환경·민주성 등 사회적 가치 지표를 중심으로 성과/파급효과를 정량·정성으로 제시합니다.",
    "경제적 성과창출의 기대효과": "당신은 재무 담당자입니다(기업 작성). 매출/원가/영업이익, ROI/NPV, 고용효과 등 재무적 성과 전망을 가정과 산식(간단) 포함하여 명료하게 제시합니다.",
    "신규 인력 채용 계획 및 활용 방안": "당신은 HR 책임자입니다. 신규/기존 채용 구분, 채용 시점·역할·배치·활용 계획, 역량 매핑과 교육/온보딩 계획을 일정표와 함께 제시합니다."
}

section_queries = {
    "연구개발 목표": "최종목표(단계/일괄 협약목표)를 과제의 연구기획목표를 500자 내외로 기재합니다.",
    "연구개발 내용": "전체내용을 1,000자 내외로 기재합니다.",
    "연구개발성과 활용계획 및 기대효과": "연구기획의 수요처, 활용내용, 경제적 파급효과 등을 500자 내외로 기재합니다(연구시설ㆍ장비 구축 과제일 경우 성과관리/자립운영/수입금 관리 계획 포함).",
    "연구기획과제의 개요": "연구개발과제의 목적, 필요성, 기대효과를 일관된 논리로 요약합니다.",
    "연구개발과제의 배경": "선행연구/시장·기술 동향/정책·RFP 부합성을 근거와 함께 기재합니다.",
    "연구개발과제의 필요성": "현황·문제점·시장규모/성장률·규제 및 정책 요구를 데이터로 제시합니다.",
    "보안등급의 분류 및 해당 사유": "법령/요령 기준을 근거로 보안등급과 결정 사유를 기재합니다.",
    "기술개발 핵심어(키워드)": "핵심 용어 5개(한글/영문) 정식 명칭으로 제시합니다.",
    "연차별 개발목표" : "연차별 목표를 기관별로 KPI/마일스톤/검증기준과 함께 정량화합니다.",
    "연차별 개발내용 및 범위" : "기관별 역할/범위/인계/의존성을 명확히 기술합니다. 명확한 기술의 예시는 가이드라인을 반드시 참고합니다. ",
    "추진방법 및 전략" : "방법론/리스크/대응책/실험·검증 계획을 구체화합니다.",
    "과제 성과의 활용방안" : "성과의 적용 분야/파급효과/환경 개선 효과 등을 기술합니다.",
    "신규사업 신설의 기대효과" : "시장/일자리/수출/비용절감 등 경제·산업적 효과를 정량 지표와 함께.",
    "사회적 가치 창출 계획" : "개요-비전-목표-세부계획-기대효과 체계로 작성합니다. ",
    "사회적 가치창출의 기대효과": "사회적 가치 지표 중심으로 정량·정성 성과를 제시합니다.",
    "경제적 성과창출의 기대효과" : "기업만 작성. 매출/원가/ROI/고용효과 등을 제시합니다.",
    "신규 인력 채용 계획 및 활용 방안" : "신규/기존 구분, 시점/배치/온보딩 계획을 포함합니다.",
}

# 6) 생성 모델 준비
gen_tok = AutoTokenizer.from_pretrained(GEN_NAME, trust_remote_code=True)
gen_tok.padding_side = "left"          
gen_tok.truncation_side = "left"        

if gen_tok.pad_token_id is None:
    gen_tok.pad_token = gen_tok.eos_token

gen_model = AutoModelForCausalLM.from_pretrained(
    GEN_NAME, trust_remote_code=True,
    torch_dtype=torch.float32 if DEVICE=="cuda" else torch.float32,
    device_map="auto"
).eval()
    
if hasattr(gen_model, "generation_config"):
    gen_model.generation_config.pad_token_id = gen_tok.pad_token_id
gen_pipe = pipeline("text-generation", model=gen_model, tokenizer=gen_tok)

# 7) 검색 + 프롬프트 구성
def build_ctx_block(section: str, project_name: str, keywords: str, topk=4):

    lines = []
    # 검색된 문서는 없으므로, combined_text만 근거로 제시
    if combined_text.strip():
        # 첫 350자만 사용 (프롬프트 길이를 위해)
        lines.append(f"- [RAG 근거] {first_n_lines(combined_text)}") 
    
    # 검색이 제거되었으므로, 쿼리 텍스트는 컨텍스트에 포함되지 않음.
    return "\n".join(lines)

def build_prompt(section: str, project_name: str, depart_name: str, project_no: str,
                 period: str, budget: str, ctx_block: str):
    role_instruction = section_roles.get(section, "당신은 연구자이자 계획서 및 해당 분야 전문가입니다.")

    request_hint   = section_queries.get(section, "문장은 ~함. ~음. 명사 등 으로 마무리하고, 각 문단마다 핵심 키워드 포함, 문장 길이/시작 다양화, 중복 표현은 피한다.")
    prompt = f"""
# #=========== 자동 문장 생성
# 역할: {role_instruction}
# 작성 항목: [{section}]
# 세부사업명: {depart_name}
# 연구개발 과제번호: {project_no}
# 연구개발과제명: {project_name}
# 전체 연구개발기간: {period}
# 예산: {budget} 천원

# 작성 조건:
#     - 제시된 {GUIDELINE_FILE} 가이드라인을 엄격히 준수하여 작성합니다.
#     - "문장은 ~함. ~음. 명사 등 으로 마무리한다. 
#     - 기술적 연관성이 낮은 미사여구는 배제합니다.
#     - 구체적인 규격/범위를 포함합니다.
#     - 아래 근거를 반영합니다:
# {ctx_block}
#     - 반드시 {RAG_JSON_FILES}의 작성 방식과 구성을 참고하여 작성합니다.
#     - 문단마다 핵심 키워드 포함, 문장 길이/시작 다양화, 중복 표현 회피.
#     - 전문적이면서 친화적인 톤.
#     - 모든 전문용어/약어에는 주석(full name) 표기.
# 요청된 형식(힌트): {request_hint}

#=========== 출력
"""
    return prompt.strip()

# (모드 1) pipeline 배치 추론
def generate_sections_pipeline_batched(prompts, batch_size=8):
    outputs = []
    bat = batch_size
    for i in tqdm(range(0, len(prompts), bat), desc="Generate (pipeline/batch)"):
        chunk = prompts[i:i+bat]
        out = gen_pipe(
            chunk,
            batch_size=bat,
            max_new_tokens=GEN_MAX_NEW_TOKENS,
            do_sample=GEN_DO_SAMPLE,
            temperature=GEN_TEMPERATURE if GEN_DO_SAMPLE else None
        )
        for item in out:
            if isinstance(item, list):
                outputs.append(item[0]["generated_text"])
            else:
                outputs.append(item["generated_text"])
    return outputs

# (모드 2) datasets + 수동 토큰화 + DataLoader + generate
def collate_pad(batch, pad_token_id):
    max_len = max(len(x["input_ids"]) for x in batch)
    input_ids = []
    attn = []
    for x in batch:
        ids = x["input_ids"]
        am = x["attention_mask"]
        pad_len = max_len - len(ids)
        input_ids.append(ids + [pad_token_id] * pad_len)
        attn.append(am + [0] * pad_len)
    return {
        "input_ids": torch.tensor(input_ids, dtype=torch.long),
        "attention_mask": torch.tensor(attn, dtype=torch.long)
    }

def generate_sections_datasets_manual(prompts, batch_size=8):
    ds = Dataset.from_dict({"prompt": prompts})

    def tok_fn(batch):
        enc = gen_tok(
            batch["prompt"],
            padding=False,
            truncation=True,
            return_attention_mask=True
        )
        return enc

    ds_tok = ds.map(tok_fn, batched=True, batch_size=1024, remove_columns=["prompt"])
    ds_tok.set_format(type=None)   

    pad_id = gen_tok.pad_token_id
    if pad_id is None:
        gen_tok.pad_token = gen_tok.eos_token
        pad_id = gen_tok.pad_token_id

    loader = DataLoader(
        ds_tok,
        batch_size=batch_size,
        shuffle=False,
        num_workers=4,
        pin_memory=True,
        persistent_workers=True,
        collate_fn=lambda b: collate_pad(b, pad_id)
    )

    model_device = next(gen_model.parameters()).device
    gen_out_texts = []

    torch.set_grad_enabled(False)
    for batch in tqdm(loader, desc="Generate (datasets/manual)"):
        input_ids = batch["input_ids"].to(model_device, non_blocking=True)
        attention_mask = batch["attention_mask"].to(model_device, non_blocking=True)

        gen_kwargs = dict(
            max_new_tokens=GEN_MAX_NEW_TOKENS,
            do_sample=GEN_DO_SAMPLE
        )
        if GEN_DO_SAMPLE and GEN_TEMPERATURE is not None:
            gen_kwargs["temperature"] = GEN_TEMPERATURE

        outputs = gen_model.generate(
            input_ids=input_ids,
            attention_mask=attention_mask,
            **gen_kwargs
        )
        
        texts = gen_tok.batch_decode(outputs, skip_special_tokens=True)
        gen_out_texts.extend(texts)

    return gen_out_texts

def generate_all_sections_batched(sections, project_name, depart_name, project_no, period, budget, keywords):
    prompts = []
    for section in sections:
        # 검색 없이 컨텍스트 블록 생성
        ctx_block = build_ctx_block(section, project_name, keywords, topk=4) 
        prompt = build_prompt(section, project_name, depart_name, project_no, period, budget, ctx_block)
        prompts.append(prompt)

    if GEN_MODE == "pipeline_batch":
        texts = generate_sections_pipeline_batched(prompts, batch_size=8)
    elif GEN_MODE == "datasets_manual":
        texts = generate_sections_datasets_manual(prompts, batch_size=8)
    else:
        raise ValueError(f"Unknown GEN_MODE: {GEN_MODE}")

    return {sec: txt for sec, txt in zip(sections, texts)}

# 8) DOCX 생성
def render_doc(project_name, depart_name, project_no, period, budget, keywords):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = '맑은 고딕'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    font.size = Pt(11)

    doc.add_heading("연구개발계획서", 0)
    doc.add_paragraph(f"세부사업명: {depart_name}")
    doc.add_paragraph(f"연구개발 과제번호: {project_no}")
    doc.add_paragraph(f"연구개발과제명: {project_name}")
    doc.add_paragraph(f"전체 연구개발기간: {period}")
    doc.add_paragraph(f"예산: {budget} 천원")
    doc.add_paragraph("")

    sections = list(section_roles.keys())

    # 검색 없이 모든 섹션에 대한 RAG 기반 텍스트 생성
    section_texts = generate_all_sections_batched(
        sections, project_name, depart_name, project_no, period, budget, keywords
    )

    for section in sections:
        doc.add_heading(section, level=1)
        
        raw_text = section_texts[section]
        cleaned_text = clean_generated_text(raw_text)
        doc.add_paragraph(cleaned_text)

    outpath = "RND_Plan_E5_No_FAISS.docx" # 파일명 변경
    doc.save(outpath)
    print(f"[DONE] {outpath} 생성 완료 (No FAISS Retrieval / GEN_MODE={GEN_MODE})")

# 9) 실행 예시
if __name__ == "__main__":
    project_name = "VUNO Med-Chest X-ray"
    depart_name  = "산업기술R&D연구기획사업"
    project_no   = "123456789"
    period       = "2023. 6. 1 ~"
    budget       = "-"
    keywords     = "흉부 엑스레이(Chest X-ray), 인공지능 의료영상(AI Medical Imaging), 폐질환 진단(Lung Disease Diagnosis), 의료 인공지능 소프트웨어(Medical AI Software), 산업기술 R&D (Industrial Technology R&D)"

    render_doc(project_name, depart_name, project_no, period, budget, keywords)