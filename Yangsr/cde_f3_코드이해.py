# ===== 라이브러리 임포트 =====
# docx: Word(.docx) 문서를 생성하고 수정하기 위한 라이브러리
from docx import Document
# transformers: Hugging Face의 사전 학습된 모델(LLM, 임베딩 모델 등)을 쉽게 사용할 수 있게 해주는 라이브러리
from transformers import AutoTokenizer, AutoModelForCausalLM, AutoModel, pipeline
# docx.shared, docx.oxml.ns: Word 문서의 글꼴 크기(Pt)와 한글 폰트(맑은 고딕) 설정을 위한 모듈
from docx.shared import Pt
from docx.oxml.ns import qn
# json: JSON 형식의 데이터를 다루기 위한 표준 라이브러리
import json
# os: 파일 시스템과 상호작용(예: 파일 경로 확인)하기 위한 라이브러리
import os
# faiss: Facebook AI에서 개발한, 대규모 벡터의 효율적인 유사성 검색 및 클러스터링을 위한 라이브러리
import faiss
# rank_bm25: BM25 알고리즘(키워드 기반 검색)을 Python에서 쉽게 사용하게 해주는 라이브러리
from rank_bm25 import BM25Okapi
# collections.defaultdict: 키가 없을 때 기본값을 자동으로 생성해주는 딕셔너리
from collections import defaultdict
# numpy: 수치 연산, 특히 배열 및 행렬 계산을 위한 핵심 라이브러리
import numpy as np
# re: 정규 표현식을 사용하여 문자열을 처리하기 위한 라이브러리
import re
# torch: PyTorch 라이브러리로, 텐서 연산 및 딥러닝 모델 구축에 사용됨
import torch

# ===== 1. 벡터 검색기: TextIndexer 클래스 =====
# 목적: 텍스트를 벡터로 변환(임베딩)하고, FAISS를 사용해 빠르게 검색할 수 있는 인덱스를 구축합니다.
#      이를 통해 의미적으로 유사한 텍스트 조각(chunk)을 효율적으로 찾을 수 있습니다. (Dense Retrieval)
class TextIndexer:
    # === 초기화 함수 (__init__) ===
    # 원인: 클래스 인스턴스 생성 시, 텍스트 임베딩에 필요한 모델과 토크나이저를 미리 로드하여 재사용성을 높이고 초기화 시간을 절약하기 위함입니다.
    # 과정: 'jxm/cde-small-v2' 모델과 해당 토크나이저를 Hugging Face에서 다운로드하여 self.tokenizer와 self.model에 저장합니다.
    # 결과: TextIndexer 객체가 생성될 때 임베딩 모델이 메모리에 로드되어 텍스트 인코딩 및 검색 준비를 마칩니다.
    def __init__(self, model_name="jxm/cde-small-v2"):
        # SentenceTransformer 대신 transformers 라이브러리를 직접 사용하여 모델과 토크나이저를 로드합니다.
        self.tokenizer = AutoTokenizer.from_pretrained(model_name, trust_remote_code=True)
        self.model = AutoModel.from_pretrained(model_name, trust_remote_code=True)
        self.idx = None  # FAISS 인덱스를 저장할 변수
        self.chunks = [] # 텍스트 조각(chunk) 리스트를 저장할 변수

    # === 텍스트 분할 함수 (chunk) ===
    # 원인: LLM이 처리하기에는 너무 긴 문서를 그대로 사용할 수 없으므로, 의미 있는 단위(섹션) 내에서 관리 가능한 길이(max_chars)로 잘라야 합니다.
    # 과정: 입력된 세그먼트(segments)를 순회하며, 각 세그먼트의 텍스트를 줄 단위로 읽어들입니다.
    #      글자 수가 max_chars를 초과하지 않도록 텍스트를 모으다가, 초과하면 하나의 청크(chunk)로 만들어 리스트(res)에 추가합니다.
    # 결과: 원본 문서가 제목(title)과 본문(text)을 포함하는 여러 개의 작은 텍스트 청크 딕셔너리 리스트로 변환됩니다. (self.chunks에 저장)
    def chunk(self, segments, max_chars=1000):
        res = []
        for seg in segments:
            buf, acc = [], 0  # buf: 텍스트 조각을 모으는 버퍼, acc: 현재 버퍼의 글자 수
            for line in seg["text"].split("\n"):
                if not line.strip(): # 빈 줄은 건너뜀
                    continue
                if acc + len(line) > max_chars and buf: # 최대 길이를 초과하고 버퍼에 내용이 있으면
                    res.append({"title": seg["title"], "text": " ".join(buf)}) # 청크 생성
                    buf, acc = [], 0 # 버퍼 초기화
                buf.append(line.strip()); acc += len(line)
            if buf: # 남아있는 버퍼 내용 처리
                res.append({"title": seg["title"], "text": " ".join(buf)})
        self.chunks = res
        return res

    # === 텍스트 인코딩(임베딩) 함수 (_encode) ===
    # 원인: 텍스트를 컴퓨터가 이해하고 유사도를 계산할 수 있는 숫자 벡터(임베딩)로 변환해야 합니다.
    # 과정: 1. 토크나이저를 사용하여 텍스트를 모델이 이해할 수 있는 숫자 ID(토큰)로 변환합니다. (padding, truncation 적용)
    #      2. 사전 학습된 언어 모델(self.model)에 토큰을 입력하여 각 토큰에 대한 벡터(last_hidden_state)를 얻습니다.
    #      3. 평균 풀링(mean pooling)을 적용하여 토큰 벡터들의 평균을 계산, 문장 전체를 대표하는 단일 벡터를 생성합니다.
    #      4. 정규화(normalize)를 통해 벡터의 길이를 1로 만들어 유사도 계산(내적)의 성능을 높입니다.
    # 결과: 입력된 텍스트 리스트에 해당하는 고차원 벡터(임베딩) 텐서가 반환됩니다.
    def _encode(self, texts, normalize=True):
        inputs = self.tokenizer(texts, padding=True, truncation=True, return_tensors="pt")
        with torch.no_grad(): # 그래디언트 계산을 비활성화하여 메모리 사용량 감소 및 계산 속도 향상
            outputs = self.model(**inputs)
        embeddings = outputs.last_hidden_state.mean(dim=1) # 평균 풀링
        if normalize:
            embeddings = torch.nn.functional.normalize(embeddings, p=2, dim=1) # L2 정규화
        return embeddings

    # === FAISS 인덱스 구축 함수 (build) ===
    # 원인: 수많은 텍스트 벡터 중에서 특정 쿼리 벡터와 가장 유사한 벡터들을 빠르게 찾기 위해, 사전 계산된 데이터 구조(인덱스)가 필요합니다.
    # 과정: 1. 'cde_minicorpus.pt' 파일 로드를 시도합니다. 이 파일에는 미리 계산된 문서 임베딩과 청크 데이터가 저장되어 있어, 매번 임베딩을 계산하는 시간을 절약합니다.
    #      2. 파일이 있다면, 임베딩 데이터를 Numpy 배열로 변환하고 FAISS의 IndexFlatIP(내적 기반 인덱스)를 생성한 뒤, 임베딩을 추가합니다.
    #      3. 파일이 없다면, self.chunks에 있는 텍스트들을 _encode 함수로 직접 임베딩하고, 그 결과로 FAISS 인덱스를 새로 구축합니다.
    # 결과: 모든 문서 청크에 대한 벡터가 포함된 FAISS 인덱스(self.idx)가 생성되어, 빠른 벡터 검색이 가능해집니다.
    def build(self, embeddings_file="cde_minicorpus.pt"):
        try:
            embeddings_data = torch.load(embeddings_file) # 미리 계산된 임베딩 파일 로드
            if isinstance(embeddings_data, dict) and 'embeddings' in embeddings_data:
                embs = embeddings_data['embeddings'].numpy().astype("float32")
                self.chunks = embeddings_data['chunks']
            elif isinstance(embeddings_data, torch.Tensor):
                embs = embeddings_data.numpy().astype("float32")
                print("Warning: Only embeddings tensor was loaded. 'chunks' data is not available.")
            else:
                raise ValueError("Unsupported file format for embeddings.")
            
            dim = embs.shape[1] # 임베딩 벡터의 차원
            self.idx = faiss.IndexFlatIP(dim) # 내적(Inner Product) 기반 인덱스 생성
            self.idx.add(embs) # 인덱스에 벡터 추가
            print("Document embeddings loaded and FAISS index built successfully from .pt file.")
        except FileNotFoundError: # 파일이 없을 경우
            print(f"Error: {embeddings_file} not found. Building embeddings from scratch.")
            texts = [c["text"] for c in self.chunks]
            if not texts:
                self.idx = None
                return
            embs = self._encode(texts).numpy().astype("float32") # 직접 임베딩
            dim = embs.shape[1]
            self.idx = faiss.IndexFlatIP(dim)
            self.idx.add(embs)

    # === 벡터 검색 수행 함수 (search_dense) ===
    # 원인: 사용자 쿼리(질문)가 주어졌을 때, FAISS 인덱스에서 의미적으로 가장 유사한 문서 청크를 찾아야 합니다.
    # 과정: 1. 'cde_query_emb.pt' 파일에서 미리 계산된 쿼리 임베딩을 로드하거나, 파일이 없으면 _encode 함수로 쿼리를 직접 임베딩하여 쿼리 벡터(qv)를 생성합니다.
    #      2. self.idx.search(qv, topk)를 호출하여 FAISS 인덱스에서 쿼리 벡터와 가장 유사한 topk개의 벡터를 찾습니다.
    #      3. 검색 결과(인덱스 I와 거리 D)를 바탕으로, 원본 청크 데이터(self.chunks)에서 해당 내용을 찾아 점수와 함께 반환합니다.
    # 결과: 쿼리와 가장 관련성이 높은 문서 청크 topk개가 담긴 리스트가 반환됩니다. 각 항목에는 텍스트, 점수, 출처("dense") 정보가 포함됩니다.
    def search_dense(self, query: str, topk=4, query_file="cde_query_emb.pt"):
        if self.idx is None or not self.chunks:
            return []
        
        try: # 미리 계산된 쿼리 임베딩 로드 시도
            query_embs = torch.load(query_file)
            if isinstance(query_embs, dict) and 'embeddings' in query_embs:
                qv = query_embs['embeddings'][0].numpy().astype("float32").reshape(1, -1)
            elif isinstance(query_embs, torch.Tensor):
                qv = query_embs.numpy().astype("float32").reshape(1, -1)
            else:
                raise ValueError("Unsupported file format for query embeddings.")
            print("Query embedding loaded successfully from .pt file.")
        except FileNotFoundError: # 파일이 없으면 직접 인코딩
            print(f"Error: {query_file} not found. Encoding query from scratch.")
            qv = self._encode([query]).numpy().astype("float32")
            
        D, I = self.idx.search(qv, topk) # FAISS 검색 실행 (D: 거리, I: 인덱스)
        items = []
        for d, i in zip(D[0], I[0]):
            if i < 0: continue # 유효하지 않은 인덱스는 건너뜀
            it = dict(self.chunks[i])
            it["score"] = float(d) # 유사도 점수
            it["source"] = "dense" # 출처: Dense 검색
            items.append(it)
        return items

# ===== 2. 하이브리드 검색기: HybridRetriever 클래스 =====
# 목적: Dense 검색(의미 기반)과 Sparse 검색(BM25, 키워드 기반)의 장점을 결합하여 검색 정확도를 높입니다.
#      예를 들어, Dense 검색이 놓칠 수 있는 특정 키워드나 약어를 BM25가 보완해줍니다.

# --- Min-Max 정규화 함수 (_minmax) ---
# 원인: 서로 다른 스케일(범위)을 가진 두 종류의 점수(Dense 점수, BM25 점수)를 공정하게 결합하려면, 두 점수를 모두 0과 1 사이의 값으로 변환(정규화)해야 합니다.
# 과정: 입력된 배열(arr)에서 최소값(lo)과 최대값(hi)을 찾고, (현재값 - 최소값) / (최대값 - 최소값) 공식을 적용합니다.
# 결과: 입력된 배열의 모든 값이 0에서 1 사이의 값으로 조정된 새로운 배열이 반환됩니다.
def _minmax(arr):
    arr = np.array(arr, dtype=float)
    if arr.size == 0: return arr
    lo, hi = float(arr.min()), float(arr.max())
    if hi - lo < 1e-12: return np.zeros_like(arr) # 분모가 0이 되는 것 방지
    return (arr - lo) / (hi - lo)

class HybridRetriever:
    # === 초기화 함수 (__init__) ===
    # 원인: 하이브리드 검색에 필요한 Dense 검색기(TextIndexer)와 BM25 모델을 초기화하고, 검색 대상이 될 전체 텍스트(corpus)를 준비해야 합니다.
    # 과정: 1. TextIndexer 인스턴스를 self.dense로 받습니다.
    #      2. TextIndexer의 청크에서 텍스트만 추출하여 self.corpus를 만듭니다.
    #      3. BM25Okapi 모델을 self.corpus를 이용해 초기화합니다. 이때 각 문서는 토크나이저(기본값: 공백 기준 분리)를 통해 단어 리스트로 변환됩니다.
    # 결과: Dense 검색과 BM25 검색을 모두 수행할 준비가 된 HybridRetriever 객체가 생성됩니다.
    def __init__(self, dense_indexer, tokenizer=lambda s: s.split(), fusion='rrf', alpha=0.6):
        self.dense = dense_indexer
        self.chunks = self.dense.chunks
        self.corpus = [c["text"] for c in self.chunks]
        self.tok = tokenizer
        self.fusion = fusion # 검색 결과 융합 방식 ('rrf' 또는 'weighted')
        self.alpha = alpha   # 가중치 융합 시 Dense 검색의 가중치
        self.bm25 = BM25Okapi([self.tok(t) for t in self.corpus]) if self.corpus else None

    # === RRF(Reciprocal Rank Fusion) 융합 함수 (_rrf) ===
    # 원인: 점수 정규화 과정 없이, 오직 '순위' 정보만을 이용해 두 검색 결과를 안정적으로 결합하고 싶을 때 사용합니다.
    # 과정: 1. Dense 검색 결과와 BM25 검색 결과 리스트를 순회하며 각 항목의 순위(r)를 매깁니다.
    #      2. 각 문서(id 기준)의 점수를 1 / (k + r) 만큼 계속 더해줍니다. (k는 순위가 점수에 미치는 영향을 조절하는 상수)
    #      3. 합산된 점수가 높은 순으로 문서를 정렬하여 최종 topk개를 선택합니다.
    # 결과: 두 검색 시스템에서 모두 상위에 나타난 문서에 더 높은 점수를 부여하는, 순위 기반의 융합된 검색 결과 리스트가 반환됩니다.
    def _rrf(self, dense_hits, bm_ranked, k=60, topk=10):
        scores = defaultdict(float)
        by_id = {} # id로 원본 아이템을 찾기 위한 맵
        for lst in [dense_hits, bm_ranked]:
            for it in lst:
                if 'id' not in it: it['id'] = id(it) # 고유 ID 부여

        for lst in [dense_hits, bm_ranked]:
            for r, it in enumerate(lst, start=1):
                _id = it.get("id")
                by_id[_id] = it
                scores[_id] += 1.0 / (k + r) # RRF 점수 계산
        
        fused = sorted(scores.items(), key=lambda x: x[1], reverse=True)[:topk]
        out = []
        for _id, s in fused:
            item = dict(by_id[_id])
            item["fused_score"] = float(s)
            item["source"] = "hybrid"
            out.append(item)
        return out

    # === 가중치 기반 융합 함수 (_weighted) ===
    # 원인: Dense 검색과 BM25 검색의 중요도를 직접 제어(alpha 값으로)하면서 점수를 결합하고 싶을 때 사용합니다.
    # 과정: 1. Dense 검색 결과의 점수와 BM25 검색 결과의 점수를 각각 _minmax 함수로 0~1 사이로 정규화합니다.
    #      2. 두 검색 결과에 나타난 모든 문서에 대해, 최종 점수 = (alpha * 정규화된_Dense점수) + ((1 - alpha) * 정규화된_BM25점수) 공식을 적용합니다.
    #      3. 계산된 최종 점수가 높은 순으로 문서를 정렬하여 topk개를 선택합니다.
    # 결과: Dense 검색과 BM25 검색의 점수가 가중 합산된, 융합된 검색 결과 리스트가 반환됩니다.
    def _weighted(self, dense_hits, bm_idx, bm_scores, topk=10):
        # Dense 결과 정규화
        d_scores = [it.get("score", 0.0) for it in dense_hits]
        d_norm = _minmax(d_scores)
        d_map = {id(it): {"item": it, "norm": float(ns)} for it, ns in zip(dense_hits, d_norm)}
        
        # BM25 결과 정규화
        b_norm = _minmax(bm_scores)
        b_map = {}
        for i, ns in zip(bm_idx, b_norm):
            it = dict(self.chunks[i])
            it['id'] = id(it); it["source"] = "bm25"
            b_map[it['id']] = {"item": it, "norm": float(ns)}
            
        merged_ids = set(d_map.keys()) | set(b_map.keys()) # 모든 문서 ID 통합
        heap = []
        for _id in merged_ids:
            dn = d_map.get(_id, {}).get("norm", 0.0)
            bn = b_map.get(_id, {}).get("norm", 0.0)
            fused = self.alpha * dn + (1.0 - self.alpha) * bn # 가중치 합산
            it = dict((d_map.get(_id) or b_map.get(_id))["item"])
            it["fused_score"] = float(fused)
            it["source"] = "hybrid"
            heap.append(it)
        
        heap.sort(key=lambda x: x["fused_score"], reverse=True)
        return heap[:topk]

    # === 하이브리드 검색 실행 함수 (search) ===
    # 원인: 사용자 쿼리에 대해 Dense 검색과 BM25 검색을 모두 실행하고, 설정된 융합 방식에 따라 최종 결과를 도출하는 통합 인터페이스가 필요합니다.
    # 과정: 1. self.dense.search_dense()를 호출하여 Dense 검색 결과를 얻습니다.
    #      2. self.bm25.get_scores()를 호출하여 모든 문서에 대한 BM25 점수를 계산하고, 점수가 높은 순으로 정렬하여 BM25 검색 결과를 얻습니다.
    #      3. self.fusion 값에 따라 _rrf 함수 또는 _weighted 함수를 호출하여 두 결과를 융합합니다.
    # 결과: 쿼리에 대해 가장 관련성 높은 문서들의 최종 융합된 리스트가 반환됩니다.
    def search(self, query, k_dense=50, k_bm25=50, topk=10, rrf_k=60):
        # 1. Dense 검색 실행
        dense_hits = self.dense.search_dense(query, topk=k_dense)
        
        # 2. BM25 검색 실행
        bm_ranked, bm_idx, bm_scores = [], [], []
        if self.bm25 is not None and self.corpus:
            qtok = self.tok(query)
            scores = self.bm25.get_scores(qtok)
            order = np.argsort(scores)[::-1][:k_bm25] # 점수가 높은 순으로 인덱스 정렬
            bm_idx = order.tolist()
            bm_scores = scores[order].tolist()
            for i in bm_idx:
                it = dict(self.chunks[i])
                it["source"] = "bm25"
                bm_ranked.append(it)
                
        # 3. 결과 융합
        if self.fusion == 'rrf':
            return self._rrf(dense_hits, bm_ranked, k=rrf_k, topk=topk)
        else:
            return self._weighted(dense_hits, bm_idx, bm_scores, topk=topk)

# ===== 3. LLM 로드 =====
# 원인: 검색된 컨텍스트(근거)와 사용자 질문을 바탕으로 최종 답변을 생성할 대규모 언어 모델(LLM)이 필요합니다.
# 과정: 1. 'skt/A.X-4.0-Light' 모델과 토크나이저를 Hugging Face에서 로드합니다.
#      2. 사용 가능한 경우 GPU(cuda)를 사용하도록 설정하여 연산 속도를 높입니다. (없으면 CPU 사용)
#      3. 로드된 모델과 토크나이저를 `pipeline`에 넣어 텍스트 생성 작업을 쉽게 수행할 수 있는 `generator`를 만듭니다.
# 결과: `generator` 객체가 생성되어, 프롬프트를 입력하면 텍스트를 생성할 준비를 마칩니다.
model_name = "skt/A.X-4.0-Light"
tokenizer = AutoTokenizer.from_pretrained(model_name)

# GPU 사용 가능 여부 확인 및 장치 설정
if torch.cuda.is_available():
    device_name = "cuda"
else:
    device_name = "cpu"
    
model = AutoModelForCausalLM.from_pretrained(
    model_name,
    dtype="auto" # 모델에 맞는 데이터 타입(예: float16)을 자동으로 설정하여 메모리 효율성 증대
).to(device_name) # 모델을 설정된 장치(GPU 또는 CPU)로 이동

# 텍스트 생성 파이프라인 설정
generator = pipeline("text-generation", model=model, tokenizer=tokenizer, device=device_name)

# ===== 4. RAG 인덱스 및 가이드라인 준비 =====
# 원인: 검색의 대상이 될 원본 문서 데이터(rag_chunks.json)와, LLM이 답변을 생성할 때 따라야 할 규칙(rnd_guideline.json)을 로드해야 합니다.
# 과정: 1. JSON_FILES 리스트에 있는 각 파일('rag_chunks.json')을 열어 텍스트 데이터를 읽어와 combined_text 변수에 모두 합칩니다.
#      2. GUIDELINE_FILE('rnd_guideline.json')을 열어 JSON 데이터를 파싱하고 guidelines 딕셔너리에 저장합니다.
#      3. 파일이 없거나 JSON 형식이 잘못된 경우, 오류 메시지를 출력하고 프로그램을 종료합니다.
# 결과: RAG 검색의 대상이 될 전체 텍스트(combined_text)와, LLM 프롬프트에 활용될 가이드라인(guidelines)이 메모리에 로드됩니다.

# --- RAG 대상 문서 로드 ---
JSON_FILES = ["rag_chunks.json"]
combined_text = ""
for json_file in JSON_FILES:
    try:
        with open(json_file, 'r', encoding='utf-8') as f:
            parsed_data = json.load(f)
            # 파일 내용 형식에 따라 텍스트 추출
            if isinstance(parsed_data, dict) and 'text' in parsed_data:
                combined_text += parsed_data['text'] + "\n\n"
            else:
                combined_text += str(parsed_data) + "\n\n"
            print(f"'{json_file}' 파일이 성공적으로 로드되었습니다.")
    except FileNotFoundError:
        print(f"오류: '{json_file}' 파일이 존재하지 않습니다. 파일을 확인해주세요.")
        exit()
    except json.JSONDecodeError:
        print(f"오류: '{json_file}' 파일의 JSON 형식이 올바르지 않습니다.")
        exit()

# --- 문서 구조화를 위한 섹션 분할 로직 ---
# 원인: 긴 문서를 의미 있는 단위(예: 'I. 서론', '1. 연구 목표')로 나누어야 검색과 내용 이해에 유리합니다.
# 과정: 정규 표현식 패턴(ANCHOR_PATTERNS)을 정의하여 문서 제목처럼 보이는 줄(앵커)을 찾습니다.
#      앵커와 앵커 사이의 텍스트를 하나의 세그먼트(segment)로 묶습니다.
# 결과: 원본 텍스트가 제목(title)과 내용(text)으로 구성된 여러 세그먼트로 분할됩니다.
ANCHOR_PATTERNS = [r"^\s*[IVXLC]+\.\s.+$", r"^\s*\d+\.\s.+$", r"^\s*\d+\-\d+\.\s.+$", r"^\s*제\s*\d+\s*기.*$", r"^\s*\(\d+\)\s.+$"]
def detect_anchors(text: str): # 텍스트에서 앵커(제목) 줄을 찾는 함수
    lines = text.splitlines()
    anchors = []
    for i, line in enumerate(lines):
        for pat in ANCHOR_PATTERNS:
            if re.match(pat, line.strip()):
                anchors.append((i, line.strip()))
                break
    return anchors

def segment_by_anchors(text: str): # 앵커를 기준으로 텍스트를 세그먼트로 나누는 함수
    lines = text.splitlines()
    anchors = detect_anchors(text)
    if not anchors: return [{"title": "FULL", "text": text}]
    segments = []
    for idx, (lineno, title) in enumerate(anchors):
        start = lineno
        end = anchors[idx+1][0] if idx+1 < len(anchors) else len(lines)
        seg_text = "\n".join(lines[start:end]).strip()
        segments.append({"title": title, "text": seg_text})
    return segments

# --- 실제 분할 및 인덱싱 실행 ---
segments = segment_by_anchors(combined_text)
indexer = TextIndexer()
chunks = indexer.chunk(segments, max_chars=1000) # 세그먼트를 다시 작은 청크로 분할
indexer.build(embeddings_file="cde_minicorpus.pt") # 청크에 대한 FAISS 인덱스 구축

# 하이브리드 검색기 인스턴스 생성
retriever = HybridRetriever(indexer, tokenizer=lambda s: s.split(), fusion='rrf', alpha=0.6)

# --- 가이드라인 파일 로드 ---
GUIDELINE_FILE = "rnd_guideline.json"
guidelines = {}
try:
    with open(GUIDELINE_FILE, 'r', encoding='utf-8') as f:
        guidelines = json.load(f)
    print(f"'{GUIDELINE_FILE}' 파일이 성공적으로 로드되었습니다.")
except FileNotFoundError:
    print(f"오류: '{GUIDELINE_FILE}' 파일이 존재하지 않습니다. 파일을 확인해주세요.")
    exit()
except json.JSONDecodeError:
    print(f"오류: '{GUIDELINE_FILE}' 파일의 JSON 형식이 올바르지 않습니다.")
    exit()

# ===== 5. 사용자 입력 및 프롬프트 엔지니어링 =====
# 원인: 사용자가 입력한 과제 정보를 바탕으로, 각 문서 섹션의 특성에 맞는 맞춤형 프롬프트를 동적으로 생성해야 LLM이 고품질의 답변을 생성할 수 있습니다.
# 과정: 1. 사용자로부터 과제 기본 정보를 입력받습니다.
#      2. `section_roles`: 각 섹션(예: "연구개발 목표")에 대해 LLM이 수행해야 할 역할(페르소나)을 정의합니다. (예: "당신은 R&D PMO입니다...")
#      3. `section_queries`: RAG 검색 시 각 섹션에 가장 적합한 검색어를 미리 정의합니다.
#      4. `build_prompt_with_context`: 검색된 근거(contexts), 역할 지침, 사용자 입력 정보, 각종 작성 규칙을 조합하여 최종 프롬프트를 구성합니다.
# 결과: LLM에게 전달될, 매우 상세하고 구조화된 지침이 담긴 프롬프트 문자열이 생성됩니다.

# --- 사용자 정보 입력 ---
depart_name = input("세부사업명: ")
project_no = input("연구개발 과제번호: ")
project_name = input("연구개발과제명: ")
period = input("전체 연구개발기간: ")
budget = input("총 연구비: ...")

# --- 섹션별 역할(페르소나) 프롬프트 정의 ---
section_roles = {
    "연구개발 목표": "당신은 R&D PMO입니다. ...",
    "연구개발 내용": "당신은 기술 총괄(Tech Lead)입니다. ...",
    # ... (나머지 섹션 정의) ...
}

# --- 섹션별 RAG 검색어 템플릿 정의 ---
keywords = "" # 이 부분은 가이드라인에서 키워드를 추출하려는 로직이지만, 실제 사용되지는 않고 있음
section_queries = {
    "연구개발 목표": "최종목표(단계/일괄 협약목표)를 과제의 연구기획목표를 500자 내외로 기재합니다.",
    "연구개발 내용": "전체내용을 1,000자 내외로 기재합니다.",
    # ... (나머지 섹션 쿼리 정의) ...
}

# ===== 6. 자동 문서 생성 로직 =====
# 원인: 위에서 준비된 모든 요소(검색기, LLM, 프롬프트)를 연결하여, 각 섹션에 대한 내용을 자동으로 생성하고 최종 문서를 완성해야 합니다.
# 과정: 1. `search_contexts`: 특정 섹션에 대해 `retriever`를 사용해 관련 문서 청크(근거)를 검색합니다.
#      2. `generate_text`: `search_contexts`로 근거를 찾고, `build_prompt_with_context`로 프롬프트를 만든 뒤, `generator`(LLM)를 호출하여 텍스트를 생성합니다.
#      3. 메인 로직: `sections` 리스트를 순회하며 각 섹션 제목을 쓰고, `generate_text`를 호출하여 해당 섹션의 본문을 생성하여 Word 문서에 추가합니다.
# 결과: 정의된 모든 섹션이 자동으로 채워진 Word(.docx) 문서가 생성됩니다.

# --- RAG 근거 검색 함수 ---
def search_contexts(section: str, topk=4):
    query = section_queries.get(section, section) # 정의된 쿼리가 있으면 사용, 없으면 섹션명을 그대로 사용
    hits = retriever.search(query, k_dense=4, k_bm25=4) # 하이브리드 검색 실행
    contexts = [h["text"][:900] + "..." if len(h["text"]) > 900 else h["text"] for h in hits] # 너무 길지 않게 자르기
    return contexts

# --- 최종 프롬프트 구성 함수 ---
def build_prompt_with_context(section, role_instruction, contexts):
    ctx_block = "\n\n".join([f"[근거]\n{c}" for c in contexts]) if contexts else "[근거]\n(해당 섹션에 대한 근거 스니펫 없음)"
    # 여러 지침과 검색된 근거, 사용자 입력을 조합하여 최종 프롬프트를 생성
    prompt = f"""
역할: {role_instruction}
작성 항목: [{section}]
세부사업명: {depart_name}
연구개발 과제번호: {project_no}

작성 조건:
    - ... (다양한 작성 규칙 명시) ...
    - {ctx_block}
    - {section} 작성 시 위 근거를 반드시 반영합니다.
    - ... (추가 규칙) ...
    """
    return prompt.strip()

# --- LLM 텍스트 생성 함수 ---
def generate_text(section, keywords=""):
    role_instruction = section_roles.get(section, "")
    contexts = search_contexts(section, topk=4) # 1. 근거 검색
    prompt = build_prompt_with_context(section, role_instruction, contexts) # 2. 프롬프트 생성
    output = generator( # 3. LLM 호출
        prompt,
        max_new_tokens=5000,
        do_sample=False, # 항상 가장 확률 높은 단어만 선택 (일관성 있는 결과)
        temperature=0,   # 생성 다양성 억제
        top_p=0.9,
        repetition_penalty=1.05, # 반복 방지
        eos_token_id=tokenizer.eos_token_id
    )
    # 생성된 텍스트에서 프롬프트 부분을 제거하고 결과만 추출
    text = output[0]["generated_text"]
    gen = text[len(prompt):].strip() if text.startswith(prompt) else text.strip()
    return gen

# ===== 7. DOCX 문서 생성 및 저장 =====
# 원인: 프로그램 실행 결과물을 사용자가 쉽게 열어보고 편집할 수 있는 표준적인 문서 형식(.docx)으로 저장해야 합니다.
# 과정: 1. `Document()`로 비어있는 Word 문서를 생성합니다.
#      2. 기본 스타일의 폰트를 '맑은 고딕', 크기 11pt로 설정하여 한글 깨짐을 방지하고 가독성을 높입니다.
#      3. 사용자 입력 정보를 문서 상단에 추가합니다.
#      4. `sections` 리스트를 순회하며, 각 섹션의 제목(add_heading)과 `generate_text`로 생성된 내용(add_paragraph)을 문서에 차례로 추가합니다.
#      5. `doc.save()`를 호출하여 완성된 문서를 파일로 저장합니다.
# 결과: '연구개발계획서(test_ver).docx' 파일이 생성되고, 프로그램이 성공적으로 완료되었음을 알리는 메시지가 출력됩니다.

doc = Document()
# --- 문서 기본 스타일(한글 폰트) 설정 ---
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
font.size = Pt(11)

# --- 문서 제목 및 기본 정보 추가 ---
doc.add_heading("연구개발계획서", 0)
doc.add_paragraph(f"세부사업명: {depart_name}")
# ... (기본 정보 추가) ...
doc.add_paragraph("")

# --- 자동 생성될 섹션 목록 ---
sections = ["연구개발 목표", "연구개발 내용", "연구개발성과 활용계획 및 기대효과", ...]

# --- 각 섹션 내용 생성 및 문서에 추가 ---
for section in sections:
    doc.add_heading(section, level=1)
    doc.add_paragraph(generate_text(section, keywords))
doc.add_page_break()

# ===== 파일 저장 =====
output_file = "연구개발계획서(test_ver).docx"
doc.save(output_file)
print(f"완료: '{output_file}' 파일이 생성되었습니다!")